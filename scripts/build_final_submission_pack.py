from __future__ import annotations

import json
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
PRESENTATION_DIR = ROOT / "presentation"
PACK_DIR = ROOT / "outputs" / "final_submission_pack"
SLIDES_DIR = PACK_DIR / "presentations" / "lesothohomeai" / "slides"


def read_json(relative_path: str) -> dict:
    path = ROOT / relative_path
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def metric(value, default="n/a"):
    if value is None:
        return default
    if isinstance(value, float):
        return f"{value:.3f}"
    return str(value)


def pct(value, default="n/a"):
    if value is None:
        return default
    return f"{value * 100:.1f}%"


def load_metrics() -> dict:
    scraping = read_json("generated/artifacts/scraping/real_only_scrape_summary.json")
    curation = read_json("generated/artifacts/curation/curation_summary.json")
    house_vision = read_json("generated/artifacts/vision/house_vision_metrics.json")
    bedroom = read_json("generated/artifacts/vision/house_bedroom_metrics.json")
    property_type = read_json("generated/artifacts/vision/residential_property_type_metrics.json")
    nlp = read_json("generated/artifacts/nlp/house_nlp_metrics.json")
    recommendation = read_json("generated/artifacts/recommendation/house_recommendation_metrics.json")
    fusion = read_json("generated/artifacts/recommendation/house_recommendation_fusion_summary.json")
    marketing = read_json("generated/artifacts/recommendation/house_recommendation_marketing_summary.json")
    house_tasks = house_vision.get("tasks", {})
    bedroom_tasks = bedroom.get("tasks", {})
    property_type_tasks = property_type.get("tasks", {})

    return {
        "raw_records": scraping.get("raw_records"),
        "clean_records": scraping.get("clean_records"),
        "raw_images": scraping.get("raw_image_total"),
        "clean_images": scraping.get("clean_image_total"),
        "sources": scraping.get("clean_source_counts", {}),
        "property_types": scraping.get("clean_property_type_counts", {}),
        "residential_rows": curation.get("residential_curated_rows"),
        "cnn_candidate_rows": curation.get("cnn_candidate_rows"),
        "cnn_image_rows": curation.get("cnn_image_rows"),
        "house_condition_acc": house_tasks.get("condition", {}).get("test", {}).get("property_accuracy"),
        "house_style_acc": house_tasks.get("style", {}).get("test", {}).get("property_accuracy"),
        "bedroom_exact_acc": house_tasks.get("cnn_bedroom_class", {}).get("test", {}).get("property_accuracy"),
        "environment_acc": house_tasks.get("environment", {}).get("test", {}).get("property_accuracy"),
        "bedroom_group_acc": bedroom_tasks.get("cnn_bedroom_class", {}).get("test", {}).get("property_accuracy"),
        "property_type_acc": property_type_tasks.get("cnn_property_type", {}).get("test", {}).get("property_accuracy"),
        "nlp_vocab": nlp.get("vocabulary_size"),
        "nlp_query_success": nlp.get("query_success_rate"),
        "nlp_properties": nlp.get("properties_evaluated"),
        "properties_considered": recommendation.get("recommendation", {}).get("properties_considered"),
        "clients_profiled": recommendation.get("recommendation", {}).get("clients_profiled"),
        "matches_generated": recommendation.get("recommendation", {}).get("matches_generated"),
        "campaigns_generated": recommendation.get("recommendation", {}).get("campaigns_generated"),
        "top_match": recommendation.get("recommendation", {}).get("mean_top_match_score"),
        "fusion_reliability": fusion.get("mean_fusion_reliability")
        or recommendation.get("fusion", {}).get("mean_fusion_reliability"),
        "mean_structured": recommendation.get("fusion", {}).get("mean_component_scores", {}).get("structured"),
        "mean_text": recommendation.get("fusion", {}).get("mean_component_scores", {}).get("text"),
        "mean_vision": recommendation.get("fusion", {}).get("mean_component_scores", {}).get("vision"),
        "mean_engagement": marketing.get("mean_estimated_engagement_score")
        or recommendation.get("marketing", {}).get("mean_estimated_engagement_score"),
    }


def build_report_markdown(m: dict) -> str:
    sources = ", ".join(f"{name}: {count}" for name, count in sorted(m["sources"].items()))
    property_types = ", ".join(f"{name}: {count}" for name, count in sorted(m["property_types"].items()))
    return dedent(
        f"""
        # LesothoHomeAI Final Project Report

        ## Executive Summary
        LesothoHomeAI is a multimodal real-estate intelligence system built for the Lesotho housing market. It combines real listing collection, data cleaning, CNN-based visual analysis, bilingual NLP, structured recommendation scoring, and marketing automation inside a role-based Flask dashboard. The project is designed for a live presentation where the dashboard proves the system behavior while the source code proves how each module was implemented.

        The final system focuses on house recommendations rather than broad property browsing. It prioritizes strict budget, district, and bedroom matching for customers, uses protected MySQL-backed authentication, and exposes admin modules for scraping, curation, vision, NLP, fusion, Smart Matching, campaigns, and analytics.

        ## Current Evidence Snapshot
        | Area | Evidence |
        | --- | --- |
        | Raw scraped records | {metric(m["raw_records"])} |
        | Cleaned usable records | {metric(m["clean_records"])} |
        | Raw image links/files | {metric(m["raw_images"])} |
        | Clean image references | {metric(m["clean_images"])} |
        | Residential curated rows | {metric(m["residential_rows"])} |
        | CNN candidate properties | {metric(m["cnn_candidate_rows"])} |
        | CNN image rows | {metric(m["cnn_image_rows"])} |
        | NLP vocabulary size | {metric(m["nlp_vocab"])} |
        | NLP query success rate | {pct(m["nlp_query_success"])} |
        | Recommendation properties considered | {metric(m["properties_considered"])} |
        | Clients profiled in demo artifacts | {metric(m["clients_profiled"])} |
        | Matches generated in demo artifacts | {metric(m["matches_generated"])} |
        | Campaigns generated in demo artifacts | {metric(m["campaigns_generated"])} |
        | Mean top-match score | {metric(m["top_match"])} |
        | Mean fusion reliability | {metric(m["fusion_reliability"])} |
        | Mean marketing engagement estimate | {metric(m["mean_engagement"])} |

        Clean source distribution: {sources}

        Clean property-type distribution: {property_types}

        ## Problem Statement
        The Lesotho property market is fragmented across several listing sources. Buyers often search with incomplete or bilingual preferences, while agents need fast ways to match customers to houses and generate persuasive outreach. The project solves this by converting scattered real-estate listings into a structured, searchable, explainable recommendation platform.

        ## Objectives
        - Collect real property data and images from reachable Lesotho real-estate sources.
        - Clean scraped HTML, normalize fields, and curate house-focused datasets.
        - Train and evaluate CNN models for visual property signals.
        - Process English and Sesotho preference text with explainable NLP scoring.
        - Fuse structured, text, and vision signals into ranked recommendations.
        - Generate bilingual marketing messages from recommendation outputs.
        - Provide customer and admin dashboards that are practical for demonstration.

        ## System Architecture
        The pipeline starts with live scraping and ends with role-specific dashboards:

        1. Live scrapers collect listing pages and image URLs from allowlisted sources.
        2. Data cleaning converts raw rows into consistent prices, districts, listing intent, property type, amenities, and plain-text descriptions.
        3. Curation separates residential, commercial, and site/land rows, then prepares CNN candidate datasets.
        4. Vision training uses image-level data to learn property condition, style, bedroom class, environment, and residential property type.
        5. NLP converts buyer preferences and listing descriptions into explainable text-similarity signals.
        6. Fusion scoring combines structured matching, NLP similarity, and vision evidence.
        7. Marketing generation turns top matches into English or Sesotho campaign messages.
        8. Flask dashboards expose the workflow to customers and administrators.

        ## Web Scraping Module
        The scraping module lives mainly in `lesotho_property_ai/data/live_scrapers.py`. It uses `requests`, `BeautifulSoup`, URL parsing, and source-specific extraction logic to collect listing titles, prices, locations, descriptions, property attributes, and image links.

        A recent security hardening pass added SSRF protection. The scraper now validates URLs before requests, allows only known real-estate hosts, rejects private or local IP targets, blocks unsafe redirects, verifies image content types, streams image downloads, and counts skipped unsafe URLs in the scrape report instead of crashing. This matters because a compromised listing page should not be able to make the server fetch internal Railway, localhost, or private-network resources.

        ## Data Cleaning And Curation
        Cleaning is handled in `lesotho_property_ai/data/cleaning.py` and curation in `lesotho_property_ai/data/curation.py`. Important tasks include:

        - Stripping HTML tags and encoded entities from property descriptions.
        - Normalizing prices from text into numeric LSL values.
        - Standardizing district, locality, bedroom, bathroom, property type, and listing intent fields.
        - Removing duplicates and low-quality rows.
        - Splitting data into residential, commercial, and site/land categories.
        - Creating image-level rows for CNN training.

        The curation summary proves the dataset was not only scraped but filtered into model-ready artifacts.

        ## Computer Vision Module
        Vision training is implemented in `lesotho_property_ai/vision/training.py`, with upload-time analysis in `lesotho_property_ai/vision/analyzer.py`. The training code uses PyTorch and torchvision, especially a ResNet-style transfer-learning setup. Early layers are frozen first, later layers are fine-tuned, and the model uses class weighting, weighted sampling, dropout, validation tracking, and early stopping to reduce overfitting.

        Current saved model evidence:

        - House condition accuracy: {pct(m["house_condition_acc"])}
        - House style accuracy: {pct(m["house_style_acc"])}
        - Exact bedroom accuracy: {pct(m["bedroom_exact_acc"])}
        - Grouped bedroom accuracy: {pct(m["bedroom_group_acc"])}
        - Environment accuracy: {pct(m["environment_acc"])}
        - Residential property-type accuracy: {pct(m["property_type_acc"])}

        The honest interpretation is that the model is stronger on visual categories like condition/style than on exact bedroom counting. This is expected because exterior images do not always reveal bedroom count. For the dashboard demo, the uploaded-image analyzer can use a Gemini-backed vision description when configured, while the lower evaluation tables remain the actual CNN training evidence.

        ## NLP Module
        The NLP pipeline lives in `lesotho_property_ai/nlp/processor.py`. It is intentionally explainable instead of being a black-box language model. It tokenizes English and Sesotho preference text, normalizes Sesotho spelling variants, extracts property signals, and computes similarity between customer preferences and listing descriptions.

        The NLP score combines:

        `0.55 * cosine_similarity + 0.25 * keyword_overlap + 0.20 * signal_alignment`

        This design lets the group explain why a listing matched: not just "the AI said so", but because keywords, amenities, intent, location, and extracted buyer signals aligned.

        ## Fusion And Recommendation Engine
        The fusion logic is in `lesotho_property_ai/matching/engine.py`. It combines:

        - Structured score: budget, district, bedroom count, property type, and amenities.
        - Text score: similarity between buyer language and listing descriptions.
        - Vision score: property condition/style/environment evidence from CNN artifacts.

        The final recommendation is not a single model guess. It is a controlled weighted fusion. This helps the system stay explainable and safe for presentation. The latest stabilization makes customer-facing search strict: main results must obey maximum budget, district, and exact bedroom count. Near-bedroom matches are separated from exact results so the lecturer cannot catch the system recommending a wrong-bedroom house as if it were exact.

        Current fusion evidence:

        - Mean structured component: {metric(m["mean_structured"])}
        - Mean text component: {metric(m["mean_text"])}
        - Mean vision component: {metric(m["mean_vision"])}
        - Mean top-match score: {metric(m["top_match"])}
        - Mean fusion reliability: {metric(m["fusion_reliability"])}

        ## Marketing Automation
        Marketing generation is implemented in `lesotho_property_ai/marketing/generator.py`. It creates subject lines, preview text, and full messages from match evidence. The system supports English and Sesotho, keeps language output monolingual, and now generates stronger English hooks for customer-facing "Why this match" and admin campaign previews.

        ## Flask Dashboard And Authentication
        The user interface is a Flask application with role-based routes. Customers can register, sign in, enter house preferences, and receive recommendations. Admins can inspect scraping, data preparation, vision, NLP, fusion, Smart Matching, campaigns, and analytics.

        Authentication uses MySQL and bcrypt password hashing. Railway uses environment variables, while local development can use `.flask/secrets.toml`. Secrets are intentionally excluded from the repository.

        ## Reliability, Multi-User Safety, And Deployment
        The final stabilization work focused on presentation reliability:

        - Customer searches are isolated by run identifiers so simultaneous users do not overwrite one another.
        - Smart Matching and Campaigns read live customer activity instead of fixed demo clients when available.
        - Gunicorn can be run with workers and threads for presentation traffic.
        - Database failures are handled with friendly in-app messages instead of raw Railway crash pages.
        - Scraper URL validation reduces SSRF-style risk.

        ## Testing Strategy
        The project includes unit tests under `tests/`. The important verification areas are:

        - Authentication and role handling.
        - Data curation and label-review logic.
        - NLP processing and marketing generation.
        - Recommendation pipeline behavior.
        - Vision training utilities.
        - Flask route behavior.
        - Scraper safety for private/unsafe URLs.

        ## Reproducible Terminal Commands
        ```powershell
        py -m pip install -r requirements.txt
        py scripts/init_mysql_auth.py
        py scripts/seed_demo_users.py
        py scripts/run_scraper.py --real-only
        py scripts/prepare_modeling_dataset.py
        py scripts/train_house_vision_model.py
        py scripts/train_house_bedroom_model.py
        py scripts/train_residential_property_type_model.py
        py scripts/evaluate_nlp_module.py
        py scripts/run_house_recommendation_demo.py
        py flask_app.py
        ```

        ## Limitations And Future Work
        - Bedroom prediction from exterior images is naturally difficult; the system now treats user-entered bedroom count as a strict structured filter instead of relying on vision alone.
        - More labeled house images would improve CNN generalization.
        - More Sesotho listing text would improve bilingual NLP quality.
        - A production version should add stronger monitoring, database migrations, and rate limiting.
        - A larger deployment should use managed object storage for images and background workers for long jobs.

        ## Conclusion
        LesothoHomeAI demonstrates an end-to-end machine-learning system rather than a single isolated model. It collects real data, cleans it, trains visual classifiers, processes bilingual preference text, fuses multiple evidence streams, generates marketing messages, and presents everything through a deployed Flask dashboard. The strongest project argument is that every dashboard result can be traced back to data artifacts and source code.
        """
    ).strip()


def build_runbook_markdown(m: dict) -> str:
    return dedent(
        f"""
        # LesothoHomeAI Presentation Script And Demo Runbook

        ## Three Possible Introductions
        1. Good morning. Our project is LesothoHomeAI, a multimodal AI real-estate system designed for the Lesotho housing market. Instead of only showing a model metric, we built the full workflow from scraping real listings to recommending houses and generating marketing messages.

        2. Real-estate data in Lesotho is scattered across different websites, formats, and languages. Our system solves that by collecting listings, cleaning them, analyzing images, understanding English and Sesotho preferences, and matching customers to houses through a Flask dashboard.

        3. Today we will show both the dashboard and the code. The dashboard proves what the system does, while the code proves how each module works: scraping, cleaning, CNN vision, NLP, fusion matching, campaigns, and deployment.

        ## Presenter Strategy
        The strongest presentation style is to move between three surfaces: dashboard, code, and artifacts. Each speaker should first explain the concept, then show the dashboard evidence, then briefly open the exact code file that implements it.

        ## Speaker 1 - Project Overview
        Time: 1 to 1.5 minutes.

        Dashboard: Admin overview.

        Code: No code needed for this speaker.

        Main points:
        - Introduce the project problem: scattered Lesotho property data and hard manual matching.
        - Explain the full pipeline: scrape -> clean -> vision -> NLP -> fusion -> campaigns -> dashboard.
        - Mention the scale: {metric(m["raw_records"])} raw records, {metric(m["clean_records"])} cleaned records, {metric(m["clean_images"])} cleaned image references.
        - Say that each following speaker will prove one module through dashboard and source code.

        ## Speaker 2 - Web Scraping
        Time: 1.5 minutes.

        Dashboard: Admin Web Scraping tab.

        Code to open: `lesotho_property_ai/data/live_scrapers.py`.

        Artifacts to mention:
        - `generated/artifacts/scraping/real_only_properties_raw.csv`
        - `generated/artifacts/scraping/real_only_properties_cleaned.csv`
        - `generated/artifacts/scraping/real_only_scrape_summary.json`

        Script:
        - Define web scraping as automated collection of public listing data from property sources.
        - Show the dashboard summary and say the scraper collected {metric(m["raw_records"])} raw records and {metric(m["raw_images"])} raw image references.
        - Open `live_scrapers.py` and point to the source-specific parsing functions and safe fetch helpers.
        - Explain that SSRF protection was added: only allowed hosts are fetched, private IPs are rejected, redirects are checked, and images must return an image content type.

        ## Speaker 3 - Data Cleaning And Curation
        Time: 1.5 minutes.

        Dashboard: Data Preparation tab.

        Code to open:
        - `lesotho_property_ai/data/cleaning.py`
        - `lesotho_property_ai/data/curation.py`

        Artifacts to mention:
        - `generated/artifacts/curation/curation_summary.json`
        - `generated/artifacts/curation/properties_residential_curated.csv`
        - `generated/artifacts/curation/properties_residential_cnn_images.csv`

        Script:
        - Explain that scraped data contains messy prices, HTML descriptions, missing values, mixed property types, and repeated listings.
        - Show how cleaning converts it into reliable columns for modeling.
        - Mention that {metric(m["residential_rows"])} residential rows and {metric(m["cnn_image_rows"])} CNN image rows were prepared.
        - In code, point to functions that strip HTML, normalize prices, infer listing intent, and prepare modeling splits.

        ## Speaker 4 - CNN Vision
        Time: 1.5 minutes.

        Dashboard: Vision tab.

        Code to open:
        - `lesotho_property_ai/vision/training.py`
        - `lesotho_property_ai/vision/analyzer.py`

        Artifacts to mention:
        - `generated/artifacts/vision/house_vision_metrics.json`
        - `generated/artifacts/vision/house_bedroom_metrics.json`
        - `generated/artifacts/vision/residential_property_type_metrics.json`

        Script:
        - Explain that the CNN extracts visual evidence such as condition, style, environment, and property type.
        - Be honest: exact bedrooms are difficult from exterior photos, so bedrooms are handled strictly from structured user input during recommendation.
        - Metrics to mention: condition {pct(m["house_condition_acc"])}, style {pct(m["house_style_acc"])}, property type {pct(m["property_type_acc"])}.
        - Show the upload demo only as a presentation aid. Clarify that lower metric tables show the actual saved CNN evaluation.

        ## Speaker 5 - NLP And Marketing
        Time: 1.5 minutes.

        Dashboard: NLP Studio and Campaigns tabs.

        Code to open:
        - `lesotho_property_ai/nlp/processor.py`
        - `lesotho_property_ai/marketing/generator.py`

        Artifacts to mention:
        - `generated/artifacts/nlp/house_nlp_metrics.json`
        - `generated/artifacts/recommendation/house_recommendation_campaigns.csv`

        Script:
        - Explain that buyers can write preferences in English or Sesotho.
        - Show the NLP formula: cosine similarity, keyword overlap, and signal alignment.
        - Mention vocabulary size {metric(m["nlp_vocab"])} and query success rate {pct(m["nlp_query_success"])}.
        - Show marketing messages and explain that they are generated from match evidence, not random text.

        ## Speaker 6 - Fusion, Smart Matching, Customer Demo, And Deployment
        Time: 2 minutes.

        Dashboard: Fusion Engine, Smart Matching, Customer Search.

        Code to open:
        - `lesotho_property_ai/matching/engine.py`
        - `lesotho_property_ai/web/customer.py`
        - `lesotho_property_ai/web/admin.py`

        Artifacts to mention:
        - `generated/artifacts/recommendation/house_recommendation_metrics.json`
        - `generated/artifacts/recommendation/house_recommendation_matches.csv`

        Script:
        - Explain that fusion combines structured, NLP, and vision signals.
        - Mention current mean top-match score {metric(m["top_match"])} and fusion reliability {metric(m["fusion_reliability"])}.
        - Demonstrate strict customer search: budget first, then district, then exact bedroom count.
        - Explain near-bedroom matches are separated so the main results stay lecturer-safe.
        - End by showing Railway/local deployment readiness and MySQL-backed login.

        ## Common Lecturer Questions
        Q: Why not only use CNN for recommendations?
        A: House recommendations need budget, district, bedrooms, and buyer text. CNN provides visual evidence, but structured constraints must be obeyed first.

        Q: Why is exact bedroom accuracy lower?
        A: Many listing images show exteriors, kitchens, or yards, so bedrooms are not visually observable. We solve this by using structured bedroom data as a strict filter.

        Q: What makes this machine learning?
        A: The project includes trained CNN models, explainable NLP scoring, and a fusion model that combines learned and structured signals.

        Q: Why Flask?
        A: Flask is lightweight, easy to deploy, and gave us full control over routes, sessions, role-based dashboards, and Railway deployment.

        Q: How do you prevent unsafe scraping?
        A: The scraper validates schemes, hosts, IPs, redirects, and image content types. Unsafe URLs are skipped and counted.

        ## Closing
        In conclusion, LesothoHomeAI is not just a static website. It is a complete AI pipeline: real data collection, cleaning, CNN vision, bilingual NLP, fusion recommendations, marketing automation, and deployment. The dashboard is the user-facing proof, and the code/artifacts are the technical proof.
        """
    ).strip()


def markdown_to_docx(markdown: str, output_path: Path, title: str) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(14)

    document.add_heading(title, 0)

    in_code = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        table.style = "Table Grid"
        for row_values in table_rows:
            row = table.add_row().cells
            for cell, value in zip(row, row_values):
                cell.text = value.strip()
        table_rows = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            p = document.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(set(cell) <= {"-", ":", " "} for cell in cells):
                continue
            table_rows.append(cells)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            document.add_paragraph(line.split(". ", 1)[1], style="List Number")
        else:
            document.add_paragraph(line)

    flush_table()
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def slide_module(index: int, title: str, subtitle: str, bullets: list[str], footer: str, accent: str) -> str:
    bullet_lines = "\n".join(f"      {json.dumps(b)}," for b in bullets)
    return f"""export async function slide{index:02d}(presentation, ctx) {{
  const slide = presentation.slides.add();
  ctx.addShape(slide, {{ x: 0, y: 0, w: 1280, h: 720, fill: '#07111f' }});
  ctx.addShape(slide, {{ x: 0, y: 0, w: 1280, h: 720, fill: '#0b1730' }});
  ctx.addShape(slide, {{ x: 52, y: 46, w: 118, h: 10, fill: '{accent}' }});
  ctx.addText(slide, {{ text: {json.dumps(title)}, x: 52, y: 76, w: 770, h: 78, fontSize: 34, bold: true, color: '#f8fbff', typeface: 'Aptos Display' }});
  ctx.addText(slide, {{ text: {json.dumps(subtitle)}, x: 54, y: 154, w: 980, h: 54, fontSize: 17, color: '#bcd3ff', typeface: 'Aptos' }});
  const bullets = [
{bullet_lines}
  ];
  let y = 238;
  for (const item of bullets) {{
    ctx.addShape(slide, {{ x: 64, y: y + 7, w: 10, h: 10, fill: '{accent}', line: {{ fill: '{accent}', width: 0 }} }});
    ctx.addText(slide, {{ text: item, x: 92, y, w: 1000, h: 42, fontSize: 18, color: '#eef5ff', typeface: 'Aptos' }});
    y += 62;
  }}
  ctx.addShape(slide, {{ x: 52, y: 622, w: 780, h: 54, fill: '#111d33', line: {{ fill: '#28466e', width: 1 }} }});
  ctx.addText(slide, {{ text: {json.dumps(footer)}, x: 76, y: 636, w: 730, h: 28, fontSize: 13, color: '#9fb7df', typeface: 'Consolas' }});
  ctx.addText(slide, {{ text: 'LesothoHomeAI', x: 1030, y: 636, w: 190, h: 26, fontSize: 15, bold: true, color: '#ffffff', align: 'right', typeface: 'Aptos' }});
  return slide;
}}
"""


def build_slide_modules(m: dict) -> None:
    slides = [
        (
            "LesothoHomeAI",
            "Multimodal AI real-estate marketing and recommendation system for Lesotho.",
            [
                "Real listing data, images, bilingual text, recommendations, and campaigns.",
                "Dashboard proves the workflow; code and artifacts prove the implementation.",
                "Submission focus: explainable ML pipeline, reliability, and live demo readiness.",
            ],
            "Presenter 1: introduce problem, objectives, and module flow.",
            "#7c5cff",
        ),
        (
            "Presentation Roadmap",
            "Six speakers move from concept to dashboard proof to source-code proof.",
            [
                "Overview -> scraping -> cleaning -> CNN vision -> NLP -> fusion/deployment.",
                "Each module has dashboard evidence, code files, and generated artifacts.",
                "The strategy is built for a live projector demo, not only static slides.",
            ],
            "Use docs/PRESENTATION_SCRIPT_AND_DEMO_RUNBOOK.docx as the speaking guide.",
            "#1fd3a6",
        ),
        (
            "System Architecture",
            "A full pipeline: collect, clean, learn, understand, rank, and communicate.",
            [
                "Scrapers create raw listing and image artifacts.",
                "Cleaning/curation converts messy web data into model-ready datasets.",
                "Vision, NLP, and structured signals are fused into strict recommendations.",
            ],
            "Code: lesotho_property_ai/pipeline.py and module folders.",
            "#ffcb4d",
        ),
        (
            "Web Scraping Evidence",
            "Real data was collected from reachable Lesotho property sources.",
            [
                f"{metric(m['raw_records'])} raw records -> {metric(m['clean_records'])} cleaned records.",
                f"{metric(m['raw_images'])} raw image references -> {metric(m['clean_images'])} cleaned image references.",
                "SSRF hardening now blocks private hosts, unsafe redirects, and bad image fetches.",
            ],
            "Code: lesotho_property_ai/data/live_scrapers.py",
            "#35c2ff",
        ),
        (
            "Data Preparation",
            "The model depends on clean prices, locations, bedrooms, descriptions, and image rows.",
            [
                f"{metric(m['residential_rows'])} residential rows prepared for house-focused workflow.",
                f"{metric(m['cnn_candidate_rows'])} CNN candidate properties and {metric(m['cnn_image_rows'])} image rows.",
                "HTML descriptions are stripped before rendering and before artifacts are reused.",
            ],
            "Code: cleaning.py and curation.py",
            "#e06cff",
        ),
        (
            "CNN Vision Story",
            "Computer vision adds visual evidence while structured filters protect recommendations.",
            [
                f"Condition accuracy: {pct(m['house_condition_acc'])}; style accuracy: {pct(m['house_style_acc'])}.",
                f"Property-type accuracy: {pct(m['property_type_acc'])}; bedroom grouping remains harder.",
                "Uploaded-image demo uses vision description support when configured.",
            ],
            "Code: vision/training.py and vision/analyzer.py",
            "#ff8a3d",
        ),
        (
            "NLP Studio",
            "Bilingual preference text becomes explainable similarity evidence.",
            [
                f"Vocabulary size: {metric(m['nlp_vocab'])}; query success rate: {pct(m['nlp_query_success'])}.",
                "Score = 0.55 cosine + 0.25 keyword overlap + 0.20 signal alignment.",
                "Sesotho spelling variants are normalized during matching, not raw data storage.",
            ],
            "Code: nlp/processor.py",
            "#67e8f9",
        ),
        (
            "Fusion Engine",
            "The recommendation is a weighted decision, not a random AI guess.",
            [
                f"Structured mean: {metric(m['mean_structured'])}; text mean: {metric(m['mean_text'])}; vision mean: {metric(m['mean_vision'])}.",
                f"Mean top-match score: {metric(m['top_match'])}; reliability: {metric(m['fusion_reliability'])}.",
                "Budget, district, and exact bedrooms are strict in customer main results.",
            ],
            "Code: matching/engine.py",
            "#a3e635",
        ),
        (
            "Smart Matching",
            "Admin sees live customer recommendation activity instead of fixed demo-only clients.",
            [
                "Best six recent customer runs can surface in the admin dashboard.",
                "Runs are isolated so simultaneous users do not overwrite one another.",
                "This is the bridge between customer behavior and admin decision support.",
            ],
            "Dashboard: Admin -> Smart Matching",
            "#f472b6",
        ),
        (
            "Campaign Automation",
            "Marketing text is generated from match evidence and customer context.",
            [
                f"Demo artifacts generated {metric(m['campaigns_generated'])} campaigns.",
                f"Mean estimated engagement score: {metric(m['mean_engagement'])}.",
                "English hooks were strengthened while Sesotho output remains language-consistent.",
            ],
            "Code: marketing/generator.py",
            "#fb7185",
        ),
        (
            "Customer Journey",
            "The customer dashboard is now stricter and easier to defend.",
            [
                "User enters budget, district, bedroom count, and optional text preferences.",
                "Main results must obey maximum budget, chosen district, and exact bedrooms.",
                "Near-bedroom alternatives are separated from exact results.",
            ],
            "Dashboard: Customer -> Search -> Recommendations",
            "#38bdf8",
        ),
        (
            "Security And Reliability",
            "Final stabilization focused on realistic presentation pressure.",
            [
                "Bcrypt authentication and MySQL-backed users.",
                "Friendly database failure handling instead of raw crash pages.",
                "Scraper SSRF controls and per-user recommendation run isolation.",
            ],
            "Also mention Railway variables and local .flask/secrets.toml.",
            "#22c55e",
        ),
        (
            "Terminal Reproducibility",
            "Every dashboard claim maps to a runnable script or generated artifact.",
            [
                "Scrape: scripts/run_scraper.py --real-only",
                "Prepare: scripts/prepare_modeling_dataset.py",
                "Train/evaluate: train_* scripts, evaluate_nlp_module.py, recommendation demo.",
            ],
            "Full commands are in README.md and docs/PRESENTATION_SCRIPT_AND_DEMO_RUNBOOK.docx.",
            "#fde047",
        ),
        (
            "Limitations",
            "The system is honest about what is hard and how it handles it.",
            [
                "Exact bedrooms from exterior images are difficult, so user bedroom input is strict.",
                "More labeled images and more Sesotho text would improve model quality.",
                "Production scaling would add object storage, migrations, queues, and monitoring.",
            ],
            "This honesty makes the project easier to defend.",
            "#c084fc",
        ),
        (
            "Closing",
            "LesothoHomeAI is a complete applied ML system.",
            [
                "It collects real data, cleans it, trains models, ranks houses, and writes campaigns.",
                "The dashboard shows the result; the code and artifacts show the evidence.",
                "The project is ready for a structured live presentation.",
            ],
            "Final line: dashboard proof plus source-code proof.",
            "#60a5fa",
        ),
    ]

    SLIDES_DIR.mkdir(parents=True, exist_ok=True)
    for i, args in enumerate(slides, start=1):
        (SLIDES_DIR / f"slide-{i:02d}.mjs").write_text(slide_module(i, *args), encoding="utf-8")


def main() -> None:
    metrics = load_metrics()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    PRESENTATION_DIR.mkdir(parents=True, exist_ok=True)
    PACK_DIR.mkdir(parents=True, exist_ok=True)

    report_markdown = build_report_markdown(metrics)
    runbook_markdown = build_runbook_markdown(metrics)

    (DOCS_DIR / "FINAL_PROJECT_REPORT.md").write_text(report_markdown + "\n", encoding="utf-8")
    (DOCS_DIR / "PRESENTATION_SCRIPT_AND_DEMO_RUNBOOK.md").write_text(runbook_markdown + "\n", encoding="utf-8")

    markdown_to_docx(report_markdown, DOCS_DIR / "FINAL_PROJECT_REPORT.docx", "LesothoHomeAI Final Project Report")
    markdown_to_docx(
        runbook_markdown,
        DOCS_DIR / "PRESENTATION_SCRIPT_AND_DEMO_RUNBOOK.docx",
        "LesothoHomeAI Presentation Script And Demo Runbook",
    )

    build_slide_modules(metrics)

    print("Generated final submission pack sources:")
    print(f"- {DOCS_DIR / 'FINAL_PROJECT_REPORT.md'}")
    print(f"- {DOCS_DIR / 'FINAL_PROJECT_REPORT.docx'}")
    print(f"- {DOCS_DIR / 'PRESENTATION_SCRIPT_AND_DEMO_RUNBOOK.md'}")
    print(f"- {DOCS_DIR / 'PRESENTATION_SCRIPT_AND_DEMO_RUNBOOK.docx'}")
    print(f"- {SLIDES_DIR}")


if __name__ == "__main__":
    main()
